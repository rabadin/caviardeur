from pathlib import Path

from docx import Document

from ..readers.base import DocumentContent


def write_docx(content: DocumentContent, output_path: Path, source_path: Path) -> None:
    """Write pseudonymized content back to a .docx file.

    Opens the original document and replaces run texts in place to preserve formatting.
    """
    doc = Document(str(source_path))

    # Build lookup: (location_type, indices) -> new text
    chunk_map: dict[tuple, str] = {}
    for chunk in content.chunks:
        loc = chunk.location
        loc_type = loc.get("type", "")

        if loc_type == "docx_run":
            key = ("docx_run", loc["para_idx"], loc["run_idx"])
            chunk_map[key] = chunk.text
        elif loc_type == "docx_table_run":
            key = (
                "docx_table_run",
                loc["table_idx"],
                loc["row_idx"],
                loc["cell_idx"],
                loc["para_idx"],
                loc["run_idx"],
            )
            chunk_map[key] = chunk.text

    # Replace paragraph runs
    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
            key = ("docx_run", para_idx, run_idx)
            if key in chunk_map:
                run.text = chunk_map[key]

    # Replace table runs
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(para.runs):
                        key = (
                            "docx_table_run",
                            table_idx,
                            row_idx,
                            cell_idx,
                            para_idx,
                            run_idx,
                        )
                        if key in chunk_map:
                            run.text = chunk_map[key]

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
