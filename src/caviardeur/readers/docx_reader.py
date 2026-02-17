from pathlib import Path

from docx import Document

from .base import DocumentContent, TextChunk


def read_docx(path: Path) -> DocumentContent:
    """Read a .docx file, extracting text at the run level to preserve formatting."""
    doc = Document(str(path))
    chunks: list[TextChunk] = []

    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
            if run.text:
                chunks.append(
                    TextChunk(
                        text=run.text,
                        location={
                            "type": "docx_run",
                            "para_idx": para_idx,
                            "run_idx": run_idx,
                        },
                    )
                )
        # Add newline between paragraphs
        chunks.append(
            TextChunk(
                text="\n",
                location={"type": "docx_separator", "para_idx": para_idx},
            )
        )

    # Also extract text from tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(para.runs):
                        if run.text:
                            chunks.append(
                                TextChunk(
                                    text=run.text,
                                    location={
                                        "type": "docx_table_run",
                                        "table_idx": table_idx,
                                        "row_idx": row_idx,
                                        "cell_idx": cell_idx,
                                        "para_idx": para_idx,
                                        "run_idx": run_idx,
                                    },
                                )
                            )
                    chunks.append(
                        TextChunk(
                            text="\n",
                            location={
                                "type": "docx_table_separator",
                                "table_idx": table_idx,
                                "row_idx": row_idx,
                                "cell_idx": cell_idx,
                            },
                        )
                    )

    content = DocumentContent(chunks=chunks, metadata={"source_path": str(path), "format": "docx"})
    content.assign_offsets()
    return content
